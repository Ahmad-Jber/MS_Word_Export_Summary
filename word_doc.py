from docx import Document

document = Document('Safety and Dependability.docx')


def add_content(head, text):

    document.add_heading(head, 1)
    document.add_paragraph(text)
    document.save('Safety and Dependability.docx')